import pandas as pd
import streamlit as st
import base64
import os
from PIL import Image
import matplotlib.pyplot as plt
import pytesseract
from docx import Document
from pdf2docx import Converter
from fpdf import FPDF
from PyPDF2 import PdfFileReader

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe' 

def converter():
    st.title("Converter arquivos")

    def get_file(bin_file, file_label='File'):
        with open(bin_file, 'rb') as f:
            data = f.read()
        bin_str = base64.b64encode(data).decode()
        href = f'<a href="data:application/octet-stream;base64,{bin_str}" download="{file_label}">Download {file_label} File</a>'
        return href

    df = None
    img = None
    file_name = None
    doc = None

    file_types = ["csv", "xls", "xlsx", "txt", "jpg", "png", "pdf", "docx"]
    uploaded_file = st.file_uploader("Arraste seu arquivo ou clique para selecionar", type=file_types)

    if uploaded_file is not None:
        file_extension = uploaded_file.name.split('.')[-1]

        if file_extension == 'csv':
            df = pd.read_csv(uploaded_file)
        elif file_extension in ['xls', 'xlsx']:
            df = pd.read_excel(uploaded_file, engine='openpyxl')
        elif file_extension == 'txt':
            df = pd.read_csv(uploaded_file, sep="\t")
        elif file_extension in ['jpg', 'png']:
            img = Image.open(uploaded_file)
        elif file_extension == 'docx':
            doc = Document(uploaded_file)
            text_data = []
            for para in doc.paragraphs:
                text_data.append(para.text.split())
            df = pd.DataFrame(text_data)
        elif file_extension == 'pdf':
            # Aqui, estou assumindo que o PDF contém texto. Se for apenas uma imagem, essa abordagem não funcionará.
            pdf_reader = PdfFileReader(uploaded_file)
            text_data = []
            for page_num in range(pdf_reader.numPages):
                text = pdf_reader.getPage(page_num).extractText()
                text_data.extend(text.split("\n"))
            df = pd.DataFrame(text_data, columns=["Text"])

        if df is not None:
            st.write('Prévia do arquivo:')
            st.dataframe(df.head())
        elif img is not None:
            st.image(img)

        st.write('---')
        col1, col2 = st.columns([2,3])
        with col1:
            option = st.selectbox('Escolha o formato para converter:', file_types)

    if st.button('Converter'):
        file_name = f"converted_file.{option}"

        if img and option in ['csv', 'xls', 'xlsx', 'txt']:
            extracted_text = pytesseract.image_to_string(img)
            data = [line.split() for line in extracted_text.split('\n') if line]
            df = pd.DataFrame(data)

        if df is not None:
            if option == 'csv':
                df.to_csv(file_name, index=False)
            elif option in ['xls', 'xlsx']:
                df.to_excel(file_name, index=False)
            elif option == 'txt':
                df.to_csv(file_name, sep="\t", index=False)
            elif option in ["jpg", "png"]:
                fig, ax = plt.subplots(figsize=(12, 4))
                ax.axis('off')
                tbl = ax.table(cellText=df.values, colLabels=df.columns, cellLoc='center', loc='center')
                tbl.auto_set_font_size(False)
                tbl.set_fontsize(9)
                for _, cell in tbl.get_celld().items():
                    cell.set_edgecolor('black')
                plt.savefig(file_name, dpi=300, bbox_inches='tight', pad_inches=0.1)
                plt.close(fig)
            elif option == "pdf":
                fig, ax = plt.subplots(figsize=(12, 4))
                ax.axis('off')
                tbl = ax.table(cellText=df.values, colLabels=df.columns, cellLoc='center', loc='center')
                tbl.auto_set_font_size(False)
                tbl.set_fontsize(9)
                for _, cell in tbl.get_celld().items():
                    cell.set_edgecolor('black')
                plt.tight_layout()
                plt.savefig(file_name, bbox_inches='tight', format='pdf')
                plt.close(fig)
            elif option == 'docx':
                doc = Document()
                for index, row in df.iterrows():
                    doc.add_paragraph('   '.join(map(str, row)))
                doc.save(file_name)
        elif file_extension == 'pdf':
            if option == 'docx':
                cv = Converter(uploaded_file.name)
                cv.convert(file_name, start=0, end=None)
                cv.close()

        st.success('Arquivo Convertido com Sucesso!')
        st.markdown(get_file(file_name, file_label=file_name), unsafe_allow_html=True)
        st.markdown("Se precisar novamente no futuro, **volte aqui** que estarei feliz em te ajudar! 😉")

        os.remove(file_name)


